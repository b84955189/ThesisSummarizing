#!/usr/bin/env python3
# -*- encoding: utf-8 -*-
"""
@File    :   WordFunc.py    
@Contact :   lking@lking.icu
@Author :    Jason
@Date :      2022/4/28 20:49
@Description  Python version-3.10

"""
import datetime
from pathlib import Path
from docx import Document

from model.RatingModels import CommentScoreModel, DebateScoreModel, TeacherScoreModel

# 模板路径
TEMPLATE_PATH = Path(__file__).parent.parent / Path("./template")
COMMENT_SCORE_FILE_PATH = TEMPLATE_PATH / Path("./comment-score.docx")
DEBATE_SCORE_FILE_PATH = TEMPLATE_PATH / Path("./debate-score.docx")
TEACHER_SCORE_FILE_PATH = TEMPLATE_PATH / Path("./teacher-score.docx")

# 替换关键字
KEY_MAJOR = "{{MAJOR}}"
KEY_STU_NUMBER = "{{SNUM}}"
KEY_STU_NAME = "{{SNAME}}"
KEY_TOPIC = "{{TOPIC}}"

# TODO: 调整分数填充方式 - 动态识别
KEY_SCORE_1 = "{{SC1}}"
KEY_SCORE_2 = "{{SC2}}"
KEY_SCORE_3 = "{{SC3}}"
KEY_SCORE_4 = "{{SC4}}"
KEY_SCORE_5 = "{{SC5}}"
KEY_SCORE_6 = "{{SC6}}"
KEY_SCORE_7 = "{{SC7}}"
KEY_SCORE_8 = "{{SC8}}"
KEY_SCORE_9 = "{{SC9}}"
KEY_SCORE_10 = "{{SC10}}"


KEY_TOTAL_SCORE = "{{TSC}}"
KEY_YEAR = "{{Y}}"
KEY_MONTH = "{{M}}"
KEY_DAY = "{{D}}"

# 年月日 - 默认当前时间
DEFAULT_YEAR = f"{datetime.datetime.now().year}"
DEFAULT_MONTH = f"{datetime.datetime.now().month}"
DEFAULT_DAY = f"{datetime.datetime.now().day}"

# 生成文件名格式 - [指导老师]-[学号]-[姓名]
COMMENT_SCORE_FILE_NAME_FORMAT = "{}-{}-{}-计科-2022届-毕业设计（论文）论文评阅评分表.docx"
DEBATE_SCORE_FILE_NAME_FORMAT = "{}-{}-{}-计科-2022届-毕业设计（论文）答辩评分表.docx"
TEACHER_SCORE_FILE_NAME_FORMAT = "{}-{}-{}-计科-2022届-毕业设计（论文）指导老师评分表.docx"

# 输出路径
OUTPUT_PATH = Path(__file__).parent.parent / Path("./out")


def generate_word_by_rating_sheet(student_number, student_name, guidance_teacher_name, rating_model, date_catalog):
    """
    通过不同评分表生成word文件（docx）
    @param student_number: 学生学号
    @param student_name: 学生姓名
    @param guidance_teacher_name: 指导老师姓名
    @param rating_model: 评分记录实体类
    @param date_catalog: 日期时间 子目录
    @return: None
    """
    output_file_path = ''
    document = ''
    if isinstance(rating_model, CommentScoreModel):
        document = Document(COMMENT_SCORE_FILE_PATH)
        # 目录
        # 创建目录
        comment_score_file_catalog_path = date_catalog / Path("./论文评阅评分")
        # parents：如果父目录不存在，是否创建父目录。
        # exist_ok：只有在目录不存在时创建目录，目录已存在时不会抛出异常。
        comment_score_file_catalog_path.mkdir(parents=True, exist_ok=True)
        output_file_path = comment_score_file_catalog_path / Path(
            COMMENT_SCORE_FILE_NAME_FORMAT.format(guidance_teacher_name,
                                                  student_number,
                                                  student_name))
    elif isinstance(rating_model, DebateScoreModel):
        document = Document(DEBATE_SCORE_FILE_PATH)
        debate_score_file_catalog_path = date_catalog / Path("./答辩评分")
        debate_score_file_catalog_path.mkdir(parents=True, exist_ok=True)
        output_file_path = debate_score_file_catalog_path / Path(
            DEBATE_SCORE_FILE_NAME_FORMAT.format(guidance_teacher_name,
                                                 student_number,
                                                 student_name))
    elif isinstance(rating_model, TeacherScoreModel):
        document = Document(TEACHER_SCORE_FILE_PATH)
        teacher_score_file_catalog_path = date_catalog / Path("./指导老师评分")
        teacher_score_file_catalog_path.mkdir(parents=True, exist_ok=True)
        output_file_path = teacher_score_file_catalog_path / Path(
            TEACHER_SCORE_FILE_NAME_FORMAT.format(guidance_teacher_name,
                                                  student_number,
                                                  student_name))

    # 因为模板中只有一个表格对象
    table = document.tables[0]
    # 遍历段落
    for paragraph in document.paragraphs:
        # 遍历run块
        for run in paragraph.runs:
            run.text = run.text.replace(KEY_MAJOR, rating_model.major)
            run.text = run.text.replace(KEY_STU_NUMBER, rating_model.student_number)
            run.text = run.text.replace(KEY_STU_NAME, rating_model.student_name)
            run.text = run.text.replace(KEY_TOPIC, rating_model.thesis_topic)

            run.text = run.text.replace(KEY_YEAR, DEFAULT_YEAR)
            run.text = run.text.replace(KEY_MONTH, DEFAULT_MONTH)
            run.text = run.text.replace(KEY_DAY, DEFAULT_DAY)
    # 遍历表格
    # 遍历 表格-行
    for row in range(len(table.rows)):
        # 遍历 表格-行的每列值
        for cell_value in table.rows[row].cells:
            # 遍历 段
            for paragraph in cell_value.paragraphs:
                # 遍历 run 块
                for run in paragraph.runs:
                    if KEY_SCORE_1 in run.text:
                        run.text = run.text.replace(KEY_SCORE_1, str(rating_model.scores[0]))
                    if KEY_SCORE_2 in run.text:
                        run.text = run.text.replace(KEY_SCORE_2, str(rating_model.scores[1]))
                    if KEY_SCORE_3 in run.text:
                        run.text = run.text.replace(KEY_SCORE_3, str(rating_model.scores[2]))
                    if KEY_SCORE_4 in run.text:
                        run.text = run.text.replace(KEY_SCORE_4, str(rating_model.scores[3]))
                    if KEY_SCORE_5 in run.text:
                        run.text = run.text.replace(KEY_SCORE_5, str(rating_model.scores[4]))
                    if KEY_SCORE_6 in run.text:
                        run.text = run.text.replace(KEY_SCORE_6, str(rating_model.scores[5]))
                    if KEY_SCORE_7 in run.text:
                        run.text = run.text.replace(KEY_SCORE_7, str(rating_model.scores[6]))
                    if KEY_SCORE_8 in run.text:
                        run.text = run.text.replace(KEY_SCORE_8, str(rating_model.scores[7]))
                    if KEY_SCORE_9 in run.text:
                        run.text = run.text.replace(KEY_SCORE_9, str(rating_model.scores[8]))
                    if KEY_SCORE_10 in run.text:
                        run.text = run.text.replace(KEY_SCORE_10, str(rating_model.scores[9]))

                    run.text = run.text.replace(KEY_TOTAL_SCORE, str(rating_model.total_score))
    # 储存新文件
    document.save(output_file_path)


def handle_output_model(output_model, date_catalog):
    """
    处理 输出实体模型
    @param output_model: 输出实体模型 - 抽象为一个学生的不同评分记录
    @param date_catalog: 日期时间 子目录
    @return:
    """
    # 根据 评阅评分记录表 生成word文件
    generate_word_by_rating_sheet(
        output_model.student_number,
        output_model.student_name,
        output_model.guidance_teacher_name,
        output_model.comment_score_model,
        date_catalog
    )
    # 根据 指导老师评分记录 生成word文件
    generate_word_by_rating_sheet(
        output_model.student_number,
        output_model.student_name,
        output_model.guidance_teacher_name,
        output_model.teacher_score_model,
        date_catalog
    )
    # 根据 答辩评分记录 生成word文件
    generate_word_by_rating_sheet(
        output_model.student_number,
        output_model.student_name,
        output_model.guidance_teacher_name,
        output_model.debate_score_model,
        date_catalog
    )


def test():
    """
    测试方法
    @return: None
    """
    template_path = Path(__file__).parent.parent / Path("./template")
    file_path = template_path / Path("./debate-score.docx")
    print(template_path)
    document = Document(file_path)
    # 因为模板中只有一个表格对象
    table = document.tables[0]
    # 遍历段落
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            # if MAJOR_KEY in run.text:
            #     print(run)
            print(run.text)
            # run.text = run.text.replace(KEY_MAJOR, "计算机科学与技术")
            # run.text = run.text.replace(KEY_STU_NAME, "刘龙龙")
    for row in range(len(table.rows)):
        # 遍历 表格-行的每列值
        for cell_value in table.rows[row].cells:
            for paragraph in cell_value.paragraphs:
                for run in paragraph.runs:
                    # if MAJOR_KEY in run.text:
                    #     print(run)
                    print(run.text)


    # 储存新文件
    document.save(template_path / Path("./new.docx"))
    # 测试创建目录
    # parents：如果父目录不存在，是否创建父目录。
    # exist_ok：只有在目录不存在时创建目录，目录已存在时不会抛出异常。
    # (template_path / Path(f'./{datetime.datetime.strftime(datetime.datetime.now(), "%Y-%m-%d-%H-%M-%S")}')).mkdir(
    #     parents=True, exist_ok=True)
    # generate_word_to_file()


if __name__ == "__main__":
    test()


